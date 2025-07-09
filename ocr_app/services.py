# ocr_app/services.py
import os
import requests
import json
import cv2
import numpy as np
import time
from datetime import datetime
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.utils import timezone
from django.conf import settings
from .models import OCRJob, OCRResult

def _upload_asset(input_data, description):
    """
    Uploads an asset to the NVCF API - exactly as per working example.
    :param input_data: The binary asset to upload
    :param description: A description of the asset
    """
    api_key = settings.NVIDIA_API_KEY
    if not api_key:
        raise ValueError("NVIDIA API key not configured")
        
    assets_url = "https://api.nvcf.nvidia.com/v2/nvcf/assets"
    header_auth = f"Bearer {api_key}"

    headers = {
        "Authorization": header_auth,
        "Content-Type": "application/json",
        "accept": "application/json",
    }

    s3_headers = {
        "x-amz-meta-nvcf-asset-description": description,
        "content-type": "image/jpeg",
    }

    payload = {"contentType": "image/jpeg", "description": description}

    response = requests.post(assets_url, headers=headers, json=payload, timeout=30)
    response.raise_for_status()

    asset_url = response.json()["uploadUrl"]
    asset_id = response.json()["assetId"]

    response = requests.put(
        asset_url,
        data=input_data,
        headers=s3_headers,
        timeout=300,
    )
    response.raise_for_status()
    return asset_id

def process_ocr_job(job_id):
    """
    Process an OCR job.
    
    Args:
        job_id (int): The ID of the OCRJob to process.
    """
    try:
        # Get the job
        job = OCRJob.objects.get(id=job_id)
        
        # Update status to processing
        job.status = OCRJob.Status.PROCESSING
        job.save()
        
        # Get the image file
        image_path = job.image.path
        start_time = time.time()
        
        # Get image dimensions
        try:
            import cv2
            image = cv2.imread(image_path)
            if image is not None:
                height, width = image.shape[:2]
                job.image_width = width
                job.image_height = height
                job.save()
        except:
            pass
        
        # Preprocess image for OCR
        try:
            preprocessed_image_path = preprocess_image_for_ocr(image_path)
            processing_image_path = preprocessed_image_path
            job.preprocessed = True
        except Exception as e:
            # Fall back to original image if preprocessing fails
            processing_image_path = image_path
            job.preprocessed = False
        
        job.save()
        
        # Perform OCR using NVIDIA API
        ocr_results = perform_nvidia_ocr(processing_image_path)
        
        # Calculate processing time
        processing_time = time.time() - start_time
        job.processing_time = processing_time
        
        # Create OCRResult objects
        create_ocr_results(job, ocr_results)
        
        # Calculate enhanced metadata
        job.total_text_items = len(ocr_results)
        if ocr_results:
            confidences = [result['confidence'] for result in ocr_results]
            job.average_confidence = sum(confidences) / len(confidences)
        
        # Generate Word document
        word_document_path = generate_word_document(job)
        
        # Update job status to completed
        job.status = OCRJob.Status.COMPLETED
        job.completed_at = timezone.now()
        job.word_document.name = word_document_path
        job.save()
        
        # Clean up preprocessed image if it was created
        if job.preprocessed and preprocessed_image_path and os.path.exists(preprocessed_image_path):
            try:
                os.unlink(preprocessed_image_path)
            except:
                pass
        
    except Exception as e:
        # Update job status to failed
        try:
            job.status = OCRJob.Status.FAILED
            job.error_message = str(e)
            job.save()
        except:
            pass
        
        # Re-raise the exception for logging
        raise

def perform_nvidia_ocr(image_path):
    """
    Perform OCR using NVIDIA's OCDRNet API with preprocessing optimizations.
    
    Args:
        image_path (str): Path to the image file.
        
    Returns:
        list: List of OCR results.
    """
    import uuid
    import zipfile
    import tempfile
    
    api_key = settings.NVIDIA_API_KEY
    
    if not api_key:
        raise ValueError("NVIDIA API key not configured")
    
    # Preprocess image for better OCR results
    preprocessed_image_path = None
    try:
        preprocessed_image_path = preprocess_image_for_ocr(image_path)
        processing_image_path = preprocessed_image_path
    except Exception as e:
        # Fall back to original image if preprocessing fails
        print(f"Image preprocessing failed: {e}")
        processing_image_path = image_path
    
    try:
        # NVAI endpoint for the ocdrnet NIM - exact from working code
        nvai_url = "https://ai.api.nvidia.com/v1/cv/nvidia/ocdrnet"
        header_auth = f"Bearer {api_key}"
        
        # Upload the asset - exactly as in working code
        with open(processing_image_path, "rb") as image_file:
            asset_id = _upload_asset(image_file.read(), "Court Document OCR")
        
        # Prepare the request with court document optimization
        inputs = {
            "image": f"{asset_id}", 
            "render_label": False  # Keep false for production use
        }
        asset_list = f"{asset_id}"
        
        headers = {
            "Content-Type": "application/json",
            "NVCF-INPUT-ASSET-REFERENCES": asset_list,
            "NVCF-FUNCTION-ASSET-IDS": asset_list,
            "Authorization": header_auth,
        }
        
        # Send request to NVIDIA OCR API - exactly as in working code
        response = requests.post(nvai_url, headers=headers, json=inputs, timeout=300)
        
        # Check for errors
        if response.status_code != 200:
            raise Exception(f"NVIDIA API returned status code {response.status_code}: {response.text}")
        
        # Handle ZIP response exactly as in the working test command
        # Create temporary files for the zip and extraction
        with tempfile.NamedTemporaryFile(suffix='.zip', delete=False) as temp_zip:
            temp_zip.write(response.content)
            temp_zip_path = temp_zip.name
        
        # Create temporary directory for extraction
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extract the zip file
            with zipfile.ZipFile(temp_zip_path, 'r') as z:
                z.extractall(temp_dir)
            
            # Find and read the .response file
            response_file = None
            for filename in os.listdir(temp_dir):
                if filename.endswith('.response'):
                    response_file = os.path.join(temp_dir, filename)
                    break
            
            if not response_file:
                raise Exception("No .response file found in the OCR result")
            
            # Read and parse the JSON response
            with open(response_file, 'r') as f:
                response_data = json.load(f)
        
        # Clean up temp zip file
        os.unlink(temp_zip_path)
        
        ocr_results = []
        
        # Enhanced processing based on model documentation insights
        if 'metadata' in response_data:
            # Filter and sort results for better court document processing
            valid_items = []
            
            for item in response_data['metadata']:
                text = item.get('label', '').strip()
                confidence = item.get('confidence', 0.0)
                polygon = item.get('polygon', {})
                
                # Apply court document specific filters
                if text and len(text) > 1:  # Skip single character noise
                    # Higher confidence threshold for court documents (0.3 instead of 0.1)
                    if confidence >= 0.3:
                        valid_items.append((item, confidence))
            
            # Sort by confidence for better quality results
            valid_items.sort(key=lambda x: x[1], reverse=True)
            
            for item, confidence in valid_items:
                text = item.get('label', '').strip()
                polygon = item.get('polygon', {})
                
                # Convert polygon format {x1, y1, x2, y2, x3, y3, x4, y4} to vertices
                vertices = []
                if polygon:
                    vertices = [
                        {'x': int(polygon.get('x1', 0)), 'y': int(polygon.get('y1', 0))},
                        {'x': int(polygon.get('x2', 0)), 'y': int(polygon.get('y2', 0))},
                        {'x': int(polygon.get('x3', 0)), 'y': int(polygon.get('y3', 0))},
                        {'x': int(polygon.get('x4', 0)), 'y': int(polygon.get('y4', 0))}
                    ]
                else:
                    vertices = [
                        {'x': 0, 'y': 0},
                        {'x': 100, 'y': 0},
                        {'x': 100, 'y': 20},
                        {'x': 0, 'y': 20}
                    ]
                
                result = {
                    'text': text,
                    'confidence': float(confidence),
                    'vertices': vertices
                }
                ocr_results.append(result)
        
        return ocr_results
            
    except json.JSONDecodeError as e:
        raise Exception(f"Failed to parse JSON response: {str(e)}")
    except Exception as e:
        raise Exception(f"Error processing OCR response: {str(e)}")
    finally:
        # Clean up preprocessed image
        if preprocessed_image_path and os.path.exists(preprocessed_image_path):
            try:
                os.unlink(preprocessed_image_path)
            except:
                pass

def format_bbox_to_vertices(bbox):
    """
    Convert bounding box to vertices format for OCR results.
    
    Args:
        bbox (dict): Bounding box coordinates
        
    Returns:
        list: List of vertex coordinates
    """
    if not bbox:
        # Default fallback
        return [
            {'x': 0, 'y': 0},
            {'x': 100, 'y': 0},
            {'x': 100, 'y': 20},
            {'x': 0, 'y': 20}
        ]
    
    # Handle different bbox formats
    if 'x' in bbox and 'y' in bbox and 'width' in bbox and 'height' in bbox:
        # Format: {x, y, width, height}
        x, y, w, h = bbox['x'], bbox['y'], bbox['width'], bbox['height']
        return [
            {'x': int(x), 'y': int(y)},
            {'x': int(x + w), 'y': int(y)},
            {'x': int(x + w), 'y': int(y + h)},
            {'x': int(x), 'y': int(y + h)}
        ]
    elif 'x1' in bbox and 'y1' in bbox and 'x2' in bbox and 'y2' in bbox:
        # Format: {x1, y1, x2, y2}
        x1, y1, x2, y2 = bbox['x1'], bbox['y1'], bbox['x2'], bbox['y2']
        return [
            {'x': int(x1), 'y': int(y1)},
            {'x': int(x2), 'y': int(y1)},
            {'x': int(x2), 'y': int(y2)},
            {'x': int(x1), 'y': int(y2)}
        ]
    elif isinstance(bbox, list) and len(bbox) >= 4:
        # Format: [x1, y1, x2, y2]
        x1, y1, x2, y2 = bbox[:4]
        return [
            {'x': int(x1), 'y': int(y1)},
            {'x': int(x2), 'y': int(y1)},
            {'x': int(x2), 'y': int(y2)},
            {'x': int(x1), 'y': int(y2)}
        ]
    
    # Default fallback
    return [
        {'x': 0, 'y': 0},
        {'x': 100, 'y': 0},
        {'x': 100, 'y': 20},
        {'x': 0, 'y': 20}
    ]

def extract_text_from_json(data):
    """
    This function is kept for compatibility but is no longer used.
    The new implementation directly parses JSON in perform_nvidia_ocr.
    """
    return []
def format_coordinates(coords):
    """
    Format coordinates to the expected vertex format.
    This function is kept for compatibility but format_bbox_to_vertices is preferred.
    """
    return format_bbox_to_vertices(coords)

def create_ocr_results(job, ocr_results):
    """
    Create OCRResult objects from the OCR results.
    
    Args:
        job (OCRJob): The job object.
        ocr_results (list): List of OCR results from NVIDIA API.
    """
    # Delete existing results if any
    job.results.all().delete()
    
    # Create new results
    for result in ocr_results:
        try:
            OCRResult.objects.create(
                job=job,
                text=result['text'],
                confidence=result['confidence'],
                x1=result['vertices'][0]['x'],
                y1=result['vertices'][0]['y'],
                x2=result['vertices'][1]['x'],
                y2=result['vertices'][1]['y'],
                x3=result['vertices'][2]['x'],
                y3=result['vertices'][2]['y'],
                x4=result['vertices'][3]['x'],
                y4=result['vertices'][3]['y']
            )
        except (KeyError, IndexError) as e:
            # Skip invalid results
            continue

def generate_word_document(job):
    """
    Generate a professional Word document from OCR results optimized for court documents.
    
    Args:
        job (OCRJob): The job object.
        
    Returns:
        str: Path to the generated Word document.
    """
    from docx.shared import RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    
    # Create a new Document
    doc = Document()
    
    # Set document margins for legal documents
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)  # Standard legal margin
        section.right_margin = Inches(1)
    
    # Add document header with legal formatting
    header_paragraph = doc.add_heading('COURT DOCUMENT - OCR TRANSCRIPTION', level=1)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add document metadata in professional format
    doc.add_paragraph('_' * 80)
    
    metadata_table = doc.add_table(rows=5, cols=2)
    metadata_table.style = 'Table Grid'
    
    metadata_cells = [
        ('Processing Date:', timezone.now().strftime("%B %d, %Y at %I:%M %p")),
        ('Document ID:', f'OCR-{job.id:06d}'),
        ('Source Image:', os.path.basename(job.image.name)),
        ('Text Items Detected:', str(job.results.count())),
        ('Processing Status:', 'COMPLETED')
    ]
    
    for i, (label, value) in enumerate(metadata_cells):
        row = metadata_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # Make labels bold
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph('_' * 80)
    
    # Add main content section
    content_heading = doc.add_heading('EXTRACTED TEXT CONTENT', level=2)
    content_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Get results ordered by reading order (top to bottom, left to right)
    results = job.results.all().order_by('y1', 'x1')
    
    # Group results by approximate line (within 15 pixels vertically)
    lines = []
    current_line = []
    last_y = None
    
    for result in results:
        if last_y is None or abs(result.y1 - last_y) <= 15:
            current_line.append(result)
        else:
            if current_line:
                # Sort current line by x-coordinate (left to right)
                current_line.sort(key=lambda r: r.x1)
                lines.append(current_line)
            current_line = [result]
        last_y = result.y1
    
    # Add the last line
    if current_line:
        current_line.sort(key=lambda r: r.x1)
        lines.append(current_line)
    
    # Add reconstructed text
    for line in lines:
        line_text = ' '.join([result.text for result in line])
        if line_text.strip():
            p = doc.add_paragraph()
            run = p.add_run(line_text)
            
            # Calculate average confidence for the line
            avg_confidence = sum([result.confidence for result in line]) / len(line)
            
            # Style based on confidence
            if avg_confidence >= 0.9:
                run.bold = True  # High confidence = bold
            elif avg_confidence < 0.5:
                run.font.color.rgb = RGBColor(128, 128, 128)  # Low confidence = gray
                run.italic = True
            
            # Add spacing between paragraphs for readability
            p.paragraph_format.space_after = Pt(6)
    
    # Add confidence analysis section
    doc.add_page_break()
    doc.add_heading('TRANSCRIPTION QUALITY ANALYSIS', level=2)
    
    # Calculate statistics
    confidences = [r.confidence for r in results]
    if confidences:
        avg_confidence = sum(confidences) / len(confidences)
        high_conf_count = len([c for c in confidences if c >= 0.8])
        medium_conf_count = len([c for c in confidences if 0.5 <= c < 0.8])
        low_conf_count = len([c for c in confidences if c < 0.5])
        
        stats_table = doc.add_table(rows=5, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ('Average Confidence:', f'{avg_confidence:.1%}'),
            ('High Confidence Items (≥80%):', f'{high_conf_count} ({high_conf_count/len(confidences):.1%})'),
            ('Medium Confidence Items (50-79%):', f'{medium_conf_count} ({medium_conf_count/len(confidences):.1%})'),
            ('Low Confidence Items (<50%):', f'{low_conf_count} ({low_conf_count/len(confidences):.1%})'),
            ('Recommendation:', 'Review low confidence items manually' if low_conf_count > 0 else 'High quality transcription')
        ]
        
        for i, (label, value) in enumerate(stats_data):
            row = stats_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True
    
    # Add detailed results table for reference
    doc.add_page_break()
    doc.add_heading('DETAILED EXTRACTION DATA', level=2)
    
    detail_table = doc.add_table(rows=1, cols=5)
    detail_table.style = 'Table Grid'
    
    # Header row
    header_cells = detail_table.rows[0].cells
    headers = ['Text Content', 'Confidence', 'Position (X,Y)', 'Dimensions', 'Quality']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for result in results:
        row_cells = detail_table.add_row().cells
        row_cells[0].text = result.text[:50] + ('...' if len(result.text) > 50 else '')
        row_cells[1].text = f'{result.confidence:.3f}'
        row_cells[2].text = f'({result.x1}, {result.y1})'
        
        # Calculate dimensions
        width = max(result.x2, result.x3) - min(result.x1, result.x4)
        height = max(result.y3, result.y4) - min(result.y1, result.y2)
        row_cells[3].text = f'{width}×{height}px'
        
        # Quality assessment
        if result.confidence >= 0.8:
            quality = 'HIGH'
        elif result.confidence >= 0.5:
            quality = 'MEDIUM'
        else:
            quality = 'LOW'
        row_cells[4].text = quality
    
    # Add footer disclaimer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        "NOTICE: This document was generated using automated OCR technology. "
        "Please review the transcription quality analysis and verify accuracy "
        "against the original document before legal use."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(10)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    document_path = f'documents/court_ocr_document_{job.id}.docx'
    full_path = os.path.join(settings.MEDIA_ROOT, document_path)
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    
    doc.save(full_path)
    
    return document_path

def preprocess_image_for_ocr(image_path):
    """
    Preprocess image for better OCR accuracy based on NVIDIA OCDRNet requirements.
    The model documentation shows it works best with grayscale images.
    
    Args:
        image_path (str): Path to the input image
        
    Returns:
        str: Path to the preprocessed image
    """
    import tempfile
    
    # Read image
    image = cv2.imread(image_path)
    if image is None:
        raise ValueError(f"Could not read image from {image_path}")
    
    # Convert to grayscale (model prefers grayscale)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    # Apply image enhancements for handwritten court documents
    # 1. Denoise the image
    denoised = cv2.fastNlMeansDenoising(gray)
    
    # 2. Enhance contrast for faded handwriting
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    enhanced = clahe.apply(denoised)
    
    # 3. Apply slight Gaussian blur to smooth out pen marks
    smoothed = cv2.GaussianBlur(enhanced, (3, 3), 0)
    
    # 4. Ensure image dimensions are multiples of 32 (as per model requirements)
    height, width = smoothed.shape
    new_height = ((height + 31) // 32) * 32
    new_width = ((width + 31) // 32) * 32
    
    # Resize if needed
    if new_height != height or new_width != width:
        smoothed = cv2.resize(smoothed, (new_width, new_height))
    
    # Save preprocessed image
    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
        cv2.imwrite(temp_file.name, smoothed)
        return temp_file.name

def detect_text_orientation(image_path):
    """
    Detect if text is rotated and return rotation angle.
    Important for court documents that might be scanned at angles.
    """
    import cv2
    import pytesseract
    
    image = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
    
    # Try different rotation angles
    angles = [0, 90, 180, 270]
    best_confidence = 0
    best_angle = 0
    
    for angle in angles:
        if angle != 0:
            # Rotate image
            center = tuple(np.array(image.shape[1::-1]) / 2)
            rot_mat = cv2.getRotationMatrix2D(center, angle, 1.0)
            rotated = cv2.warpAffine(image, rot_mat, image.shape[1::-1], flags=cv2.INTER_LINEAR)
        else:
            rotated = image
        
        # Use Tesseract to get confidence
        try:
            data = pytesseract.image_to_data(rotated, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            if avg_confidence > best_confidence:
                best_confidence = avg_confidence
                best_angle = angle
        except:
            continue
    
    return best_angle
